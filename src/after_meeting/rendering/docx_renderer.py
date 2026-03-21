"""DOCX renderer — generates professional Korean-language meeting documents."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from after_meeting.models import StructuredMeeting
from after_meeting.rendering import register


class DocxRenderer:
    """Render a StructuredMeeting to a .docx file programmatically."""

    # ------------------------------------------------------------------
    # Style constants
    # ------------------------------------------------------------------
    _TITLE_SIZE = Pt(16)
    _HEADING_SIZE = Pt(14)
    _BODY_SIZE = Pt(11)
    _FONT_NAME = "맑은 고딕"
    _HEADER_BG = RGBColor(0xD9, 0xD9, 0xD9)  # light gray

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def render(self, meeting: StructuredMeeting, output_path: Path) -> Path:
        """Create a .docx document from *meeting* and write it to *output_path*."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._set_default_font(doc)

        if meeting.doc_type == "report":
            self._build_report(doc, meeting)
        else:
            self._build_minutes(doc, meeting)

        # Appendix: full transcript (both types)
        if meeting.full_transcript is not None:
            self._add_transcript_appendix(doc, meeting)

        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Minutes (회의록)
    # ------------------------------------------------------------------
    def _build_minutes(self, doc: Document, meeting: StructuredMeeting) -> None:
        self._add_title(doc, "회의록")

        # 1. 회의 정보
        self._add_section_heading(doc, "1. 회의 정보")
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "회의 제목"
        table.rows[0].cells[1].text = "회의 일자"
        self._style_header_row(table.rows[0])
        table.rows[1].cells[0].text = meeting.title
        table.rows[1].cells[1].text = meeting.date
        self._set_table_body_font(table, start_row=1)
        doc.add_paragraph()

        # 2. 안건별 논의 내용
        self._add_section_heading(doc, "2. 안건별 논의 내용")
        for idx, agenda in enumerate(meeting.agenda_discussions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {agenda.topic}")
            run.bold = True
            run.font.size = Pt(12)
            self._set_run_font(run)

            summary_p = doc.add_paragraph(agenda.summary)
            self._set_paragraph_font(summary_p)

            if agenda.speaker_contributions:
                for contrib in agenda.speaker_contributions:
                    speaker = contrib.get("speaker", "")
                    content = contrib.get("contribution", "")
                    cp = doc.add_paragraph(style="List Bullet")
                    run = cp.add_run(f"{speaker}: ")
                    run.bold = True
                    self._set_run_font(run)
                    run2 = cp.add_run(content)
                    self._set_run_font(run2)
            doc.add_paragraph()

        # 3. 결정사항
        self._add_section_heading(doc, "3. 결정사항")
        self._add_numbered_list(doc, meeting.decisions)
        doc.add_paragraph()

        # 4. 액션 아이템
        self._add_section_heading(doc, "4. 액션 아이템")
        self._add_action_items_table(doc, meeting)

    # ------------------------------------------------------------------
    # Report (회의보고서)
    # ------------------------------------------------------------------
    def _build_report(self, doc: Document, meeting: StructuredMeeting) -> None:
        self._add_title(doc, "회의보고서")

        # 1. 회의 개요
        self._add_section_heading(doc, "1. 회의 개요")
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "회의 제목"
        table.rows[0].cells[1].text = "회의 일자"
        self._style_header_row(table.rows[0])
        table.rows[1].cells[0].text = meeting.title
        table.rows[1].cells[1].text = meeting.date
        self._set_table_body_font(table, start_row=1)
        doc.add_paragraph()

        # 2. 주요 논의 요약
        self._add_section_heading(doc, "2. 주요 논의 요약")
        if meeting.executive_summary:
            p = doc.add_paragraph(meeting.executive_summary)
            self._set_paragraph_font(p)
            doc.add_paragraph()

        for idx, agenda in enumerate(meeting.agenda_discussions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {agenda.topic}")
            run.bold = True
            run.font.size = Pt(12)
            self._set_run_font(run)

            summary_p = doc.add_paragraph(agenda.summary)
            self._set_paragraph_font(summary_p)
        doc.add_paragraph()

        # 3. 결정사항
        self._add_section_heading(doc, "3. 결정사항")
        self._add_numbered_list(doc, meeting.decisions)
        doc.add_paragraph()

        # 4. 후속조치 계획
        self._add_section_heading(doc, "4. 후속조치 계획")
        self._add_action_items_table(doc, meeting)

    # ------------------------------------------------------------------
    # Appendix: Full Transcript
    # ------------------------------------------------------------------
    def _add_transcript_appendix(
        self, doc: Document, meeting: StructuredMeeting
    ) -> None:
        doc.add_page_break()
        self._add_section_heading(doc, "부록: 전체 회의 기록")

        transcript = meeting.full_transcript
        assert transcript is not None  # caller guarantees this

        for utt in transcript.utterances:
            ts = self._format_timestamp(utt.start_time)
            p = doc.add_paragraph()
            ts_run = p.add_run(f"[{ts}] ")
            ts_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            ts_run.font.size = self._BODY_SIZE
            self._set_run_font(ts_run)

            speaker_run = p.add_run(f"{utt.speaker}: ")
            speaker_run.bold = True
            speaker_run.font.size = self._BODY_SIZE
            self._set_run_font(speaker_run)

            text_run = p.add_run(utt.text)
            text_run.font.size = self._BODY_SIZE
            self._set_run_font(text_run)

    # ------------------------------------------------------------------
    # Shared helpers
    # ------------------------------------------------------------------
    def _set_default_font(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = self._FONT_NAME
        font.size = self._BODY_SIZE

    def _add_title(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = self._TITLE_SIZE
        self._set_run_font(run)
        doc.add_paragraph()  # spacer

    def _add_section_heading(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = self._HEADING_SIZE
        self._set_run_font(run)

    def _add_numbered_list(self, doc: Document, items: list[str]) -> None:
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph(f"{idx}. {item}")
            self._set_paragraph_font(p)

    def _add_action_items_table(
        self, doc: Document, meeting: StructuredMeeting
    ) -> None:
        n_items = len(meeting.action_items)
        table = doc.add_table(rows=1 + n_items, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        header_row = table.rows[0]
        header_row.cells[0].text = "담당자"
        header_row.cells[1].text = "내용"
        header_row.cells[2].text = "기한"
        self._style_header_row(header_row)

        # Column widths (approximate)
        for row in table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(10)
            row.cells[2].width = Cm(3)

        # Data
        for i, item in enumerate(meeting.action_items):
            row = table.rows[i + 1]
            row.cells[0].text = item.assignee or "-"
            row.cells[1].text = item.description
            row.cells[2].text = item.deadline or "-"

        self._set_table_body_font(table, start_row=1)

    def _style_header_row(self, row) -> None:
        """Apply bold + light gray background to every cell in *row*."""
        from docx.oxml.ns import qn

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = self._BODY_SIZE
                    self._set_run_font(run)

            # Background shading
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            shading = tc_pr.makeelement(
                qn("w:shd"),
                {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "D9D9D9",
                },
            )
            tc_pr.append(shading)


    def _set_table_body_font(self, table, *, start_row: int = 0) -> None:
        for row in table.rows[start_row:]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = self._BODY_SIZE
                        self._set_run_font(run)

    def _set_paragraph_font(self, paragraph) -> None:
        for run in paragraph.runs:
            run.font.size = self._BODY_SIZE
            self._set_run_font(run)

    def _set_run_font(self, run) -> None:
        run.font.name = self._FONT_NAME

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Convert seconds to HH:MM:SS."""
        total = int(seconds)
        h, remainder = divmod(total, 3600)
        m, s = divmod(remainder, 60)
        return f"{h:02d}:{m:02d}:{s:02d}"


# Register with the rendering registry
register("docx", DocxRenderer)
