"""Tests for the DOCX renderer."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from after_meeting.models import (
    ActionItem,
    AgendaDiscussion,
    StructuredMeeting,
    Transcript,
    Utterance,
)
from after_meeting.rendering.docx_renderer import DocxRenderer

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "fixtures"


# ------------------------------------------------------------------
# Fixtures
# ------------------------------------------------------------------
@pytest.fixture()
def sample_meeting() -> StructuredMeeting:
    """Load the canonical sample meeting from the JSON fixture."""
    with open(FIXTURES_DIR / "sample_structured.json", encoding="utf-8") as f:
        data = json.load(f)
    return StructuredMeeting(**data)


@pytest.fixture()
def sample_meeting_report(sample_meeting: StructuredMeeting) -> StructuredMeeting:
    """A report variant of the sample meeting."""
    return sample_meeting.model_copy(
        update={
            "doc_type": "report",
            "executive_summary": (
                "본 회의에서는 Q2 출시 일정과 사용자 피드백 반영 계획을 "
                "논의하였으며, 주요 결정사항과 액션 아이템을 도출하였습니다."
            ),
        }
    )


@pytest.fixture()
def sample_meeting_with_transcript(
    sample_meeting: StructuredMeeting,
) -> StructuredMeeting:
    """Sample meeting with a full transcript attached."""
    transcript = Transcript(
        language="ko",
        speakers=["Speaker 1", "Speaker 2"],
        utterances=[
            Utterance(
                speaker="Speaker 1",
                start_time=0.0,
                end_time=5.5,
                text="그럼 회의를 시작하겠습니다.",
            ),
            Utterance(
                speaker="Speaker 2",
                start_time=6.0,
                end_time=12.3,
                text="네, 준비되었습니다.",
            ),
            Utterance(
                speaker="Speaker 1",
                start_time=13.0,
                end_time=25.0,
                text="먼저 Q2 출시 일정부터 검토하겠습니다.",
            ),
        ],
    )
    return sample_meeting.model_copy(update={"full_transcript": transcript})


@pytest.fixture()
def renderer() -> DocxRenderer:
    return DocxRenderer()


# ------------------------------------------------------------------
# Tests
# ------------------------------------------------------------------
class TestDocxRendererMinutes:
    """Tests for minutes (회의록) rendering."""

    def test_produces_valid_docx(
        self, renderer: DocxRenderer, sample_meeting: StructuredMeeting, tmp_path: Path
    ) -> None:
        out = tmp_path / "minutes.docx"
        result = renderer.render(sample_meeting, out)

        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_output_path_is_returned(
        self, renderer: DocxRenderer, sample_meeting: StructuredMeeting, tmp_path: Path
    ) -> None:
        out = tmp_path / "subdir" / "minutes.docx"
        result = renderer.render(sample_meeting, out)
        assert result == out
        assert out.exists()

    def test_docx_contains_expected_content(
        self, renderer: DocxRenderer, sample_meeting: StructuredMeeting, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "minutes.docx"
        renderer.render(sample_meeting, out)
        doc = Document(str(out))

        para_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        full_text = para_text + "\n" + table_text

        assert "회의록" in full_text
        assert "회의 정보" in full_text
        assert "안건별 논의 내용" in full_text
        assert "결정사항" in full_text
        assert "액션 아이템" in full_text
        assert sample_meeting.title in full_text

    def test_no_appendix_without_transcript(
        self, renderer: DocxRenderer, sample_meeting: StructuredMeeting, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "minutes.docx"
        renderer.render(sample_meeting, out)
        doc = Document(str(out))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "부록" not in full_text


class TestDocxRendererReport:
    """Tests for report (회의보고서) rendering."""

    def test_produces_valid_docx(
        self,
        renderer: DocxRenderer,
        sample_meeting_report: StructuredMeeting,
        tmp_path: Path,
    ) -> None:
        out = tmp_path / "report.docx"
        result = renderer.render(sample_meeting_report, out)

        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_docx_contains_report_sections(
        self,
        renderer: DocxRenderer,
        sample_meeting_report: StructuredMeeting,
        tmp_path: Path,
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        renderer.render(sample_meeting_report, out)
        doc = Document(str(out))

        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "회의보고서" in full_text
        assert "회의 개요" in full_text
        assert "주요 논의 요약" in full_text
        assert "결정사항" in full_text
        assert "후속조치 계획" in full_text

    def test_executive_summary_included(
        self,
        renderer: DocxRenderer,
        sample_meeting_report: StructuredMeeting,
        tmp_path: Path,
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        renderer.render(sample_meeting_report, out)
        doc = Document(str(out))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert sample_meeting_report.executive_summary in full_text


class TestDocxRendererWithTranscript:
    """Tests for the optional transcript appendix."""

    def test_appendix_present(
        self,
        renderer: DocxRenderer,
        sample_meeting_with_transcript: StructuredMeeting,
        tmp_path: Path,
    ) -> None:
        from docx import Document

        out = tmp_path / "minutes_with_appendix.docx"
        renderer.render(sample_meeting_with_transcript, out)
        doc = Document(str(out))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "부록: 전체 회의 기록" in full_text

    def test_appendix_contains_utterances(
        self,
        renderer: DocxRenderer,
        sample_meeting_with_transcript: StructuredMeeting,
        tmp_path: Path,
    ) -> None:
        from docx import Document

        out = tmp_path / "minutes_with_appendix.docx"
        renderer.render(sample_meeting_with_transcript, out)
        doc = Document(str(out))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Check timestamps and speaker names appear
        assert "[00:00:00]" in full_text
        assert "Speaker 1:" in full_text
        assert "Speaker 2:" in full_text
        assert "그럼 회의를 시작하겠습니다." in full_text

    def test_report_with_transcript_appendix(
        self,
        renderer: DocxRenderer,
        sample_meeting_with_transcript: StructuredMeeting,
        tmp_path: Path,
    ) -> None:
        """Report doc_type should also get the appendix when transcript exists."""
        from docx import Document

        meeting = sample_meeting_with_transcript.model_copy(
            update={"doc_type": "report"}
        )
        out = tmp_path / "report_with_appendix.docx"
        renderer.render(meeting, out)
        doc = Document(str(out))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "회의보고서" in full_text
        assert "부록: 전체 회의 기록" in full_text


class TestDocxRendererRegistry:
    """Ensure the renderer is properly registered."""

    def test_registered_as_docx(self) -> None:
        from after_meeting.rendering import get_renderer

        r = get_renderer("docx")
        assert isinstance(r, DocxRenderer)
